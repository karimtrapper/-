# -*- coding: utf-8 -*-
"""Генерация пакета документов MF Corp → клиент по проверенным юристом шаблонам.

Три флоу — три шаблона: freehold / leasehold / rental. Договор рамочный,
подписывается один раз на клиента И НА ТИП СДЕЛКИ (шаблоны юридически разные),
каждый последующий платёж оформляется своим Приложением 1 плюс инвойсом.

Правила зашиты по карте полей (wiki `crm-doc-generator-fields.md`):
  * комиссия по умолчанию вшита в курс, отдельно не взимается;
  * OUR/SHA/BEN — только фрихолд со SWIFT, на лизхолде и аренде внутренний
    тайский перевод без допрасходов;
  * назначение платежа генерируется и зависит от способа pay-in: на СБП его нет;
  * строки таблиц ищем по фрагменту текста, а не по индексу — юрист добавляет
    пункты и нумерация едет.
"""
from __future__ import annotations

import copy
import io
import os
import re
from datetime import datetime

TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'doc_templates')
ASSETS = os.path.join(TEMPLATE_DIR, 'assets')

TEMPLATES = {
    'freehold': 'MF_Freehold_Payment_Agreement_Template_RU_EN.docx',
    'leasehold': 'MF_Leasehold_Payment_Agreement_Template_RU_EN.docx',
    'rental': 'MF_Rental_Payment_Agreement_Template_RU_EN_v2.docx',
}
INVOICE_TEMPLATE = 'MF_Invoice_Short_Template_RU_EN_v2.docx'
# Рублёвый «Коммерческий инвойс» — то, что реально уходит клиенту на оплату.
# Формат снят с MF-180-2808-1: позиция, назначение капсом, реквизиты ООО.
COMMERCIAL_INVOICE_TEMPLATE = 'MF_Commercial_Invoice_RU.docx'
PEC_TEMPLATE = 'MF_Payment_Execution_Confirmation_Template_RU_EN.docx'

DEAL_TYPE_TITLES = {
    'freehold': 'Фрихолд — покупка в собственность',
    'leasehold': 'Лизхолд — покупка права аренды',
    'rental': 'Аренда — депозит, арендная плата, коммунальные',
}

AGENT = {
    'name': 'MF Corporation Company Limited',
    'reg': '0835565024547',
    'address': '9/31, Mu 5, Choeng Thale Sub-district, Thalang District, Phuket Province 83110, Thailand',
    'director': 'Miss Katika Sakornnoi',
    'email': 'mfcorpthai@gmail.com',
}

MONTHS_RU = ['января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
             'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря']
MONTHS_EN = ['January', 'February', 'March', 'April', 'May', 'June',
             'July', 'August', 'September', 'October', 'November', 'December']

# Дефолты, снятые с живого договора Фролова (аренда, 01.09.2026)
# Срок исполнения различается по флоу — сверено с живыми договорами:
# аренда Фролова 1 рабочий день, фрихолд Антоненко и лизхолд Буровой по 3.
EXECUTION_DAYS = {'freehold': '3', 'leasehold': '3', 'rental': '1'}

DEFAULTS = {
    'execution_days': '3',
    'report_days': '5',
    'terminate_days': '10',
    'remaining_balance': 'отсутствует / none',
    'confirmation_contact': f"{AGENT['email']}; менеджер MF Corporation / MF Corporation manager",
    'bank_charges_local': ('Внутренний тайский перевод; дополнительные банковские расходы '
                           'с Клиента не взимаются / Domestic Thai transfer; no additional '
                           'bank charges are payable by the Client'),
    'fee_included': 'Включена в курс и в итоговую сумму / Included in the rate and in the total',
}

PAYIN_METHODS = {
    'bank': 'Банковский перевод / Bank transfer',
    'sbp': 'СБП (Система быстрых платежей) / SBP (Faster Payments System)',
    'usdt': 'USDT',
    'cash': 'Наличные / Cash',
}
PAYIN_EVIDENCE = {
    'bank': 'Платёжное поручение / Payment order',
    'sbp': 'Чек банковского приложения / Bank app receipt',
    'usdt': 'TXID транзакции / Transaction TXID',
    'cash': 'Кассовый документ / Cash receipt',
}


# ─────────────────────────── нумерация ───────────────────────────

def make_number(passport_no: str, when: datetime | None = None, seq: int = 1) -> str:
    """`MF-<3 последние цифры паспорта>-<ДДММ>-<номер соглашения>`.

    Схема Карима от 03.09.2026, подтверждена на договоре Бранова
    (паспорт 77 6892733, 16.06 → MF-733-1606).
    """
    digits = re.sub(r'\D', '', passport_no or '')
    tail = (digits[-3:] or '000').rjust(3, '0')
    when = when or datetime.now()
    return f'MF-{tail}-{when:%d%m}-{seq}'


def date_ru_en(when: datetime) -> str:
    return (f'{when.day} {MONTHS_RU[when.month - 1]} {when.year} / '
            f'{when.day} {MONTHS_EN[when.month - 1]} {when.year}')


# ─────────────────────── низкоуровневые правки docx ───────────────────────

def _para_replace(p, old: str, new: str) -> bool:
    """Замена подстроки в параграфе со склейкой runs (формат первого run)."""
    full = ''.join(r.text for r in p.runs)
    if old not in full:
        return False
    full = full.replace(old, new)
    for r in p.runs[1:]:
        r.text = ''
    if p.runs:
        p.runs[0].text = full
    return True


def _set_cell(cell, text: str) -> None:
    """Переписать ячейку, сохранив формат первого run. \n → перенос строки."""
    lines = str(text).split('\n')
    p = cell.paragraphs[0]
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    if not p.runs:
        p.add_run('')
    base = p.runs[0]
    for r in p.runs[1:]:
        r._element.getparent().remove(r._element)
    base.text = lines[0]
    for line in lines[1:]:
        nr = copy.deepcopy(base._element)
        base._element.addnext(nr)
        base = p.runs[-1]
        base.text = line
        base._element.insert(0, base._element.makeelement(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br', {}))


def _find_row(table, needle: str):
    """Строка по фрагменту текста — устойчиво к сдвигу нумерации пунктов."""
    for row in table.rows:
        if any(needle in c.text for c in row.cells):
            return row
    return None


def _set_field(table, label: str, value) -> bool:
    """Приложения 1 и 2 — таблицы «метка | значение», значение в последней ячейке."""
    row = _find_row(table, label)
    if row is None or value in (None, ''):
        return False
    _set_cell(row.cells[-1], value)
    return True


def _unique_cells(doc):
    """Все ячейки документа по одному разу.

    Объединённая ячейка возвращается из row.cells несколько раз. Сравниваем сами
    элементы <w:tc>, а НЕ id() от них: lxml пересоздаёт прокси, адреса
    переиспользуются после сборки мусора, и дедупликация по id молча ломается.
    """
    cells, seen = [], []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if any(cell._tc is t for t in seen):
                    continue
                seen.append(cell._tc)
                cells.append(cell)
    return cells


def _sign_agent(doc, when: datetime) -> None:
    """Подпись Катики, дата и печать в блоки АГЕНТА — сразу, без прочерков.

    Клиенту документ уходит уже подписанным с нашей стороны: так просил Карим,
    прочерков под подпись Агента не оставляем.
    """
    from docx.shared import Inches  # noqa: PLC0415

    sig = os.path.join(ASSETS, 'signature.png')
    stamp = os.path.join(ASSETS, 'stamp.png')
    date_txt = f'Date: {when:%d.%m.%Y} / {when.day} {MONTHS_EN[when.month - 1]} {when.year}'
    for cell in _unique_cells(doc):
        text = cell.text
        # блок АГЕНТА: и в теле договора, и в приложении. У клиента в той же
        # строке отдельная ячейка — её не трогаем, там подпись ставит клиент.
        if AGENT['director'] not in text and 'Registration / Tax No.' not in text:
            continue
        signed = False
        for p in cell.paragraphs:
            for r in p.runs:
                if r.text.strip().startswith('Signature:') and os.path.exists(sig):
                    r.text = 'Signature: '
                    r.add_picture(sig, width=Inches(1.4))
                    signed = True
                elif r.text.strip().startswith('Date:'):
                    r.text = date_txt
        if signed and os.path.exists(stamp):
            sp = cell.add_paragraph()
            sp.add_run().add_picture(stamp, width=Inches(1.2))


def _strip_draft_mark(doc) -> None:
    for p in list(doc.paragraphs):
        if p.text.strip() == 'ОБЕЗЛИЧЕННЫЙ ШАБЛОН':
            p._element.getparent().remove(p._element)


# ─────────────────────── сборка значений полей ───────────────────────

# Гражданство в паспорте написано по-русски, а в английской колонке договора
# кириллица выглядит браком. Переводим известные, остальное оставляем как есть.
CITIZENSHIP_EN = {
    'российской федерации': 'the Russian Federation', 'россии': 'the Russian Federation',
    'рф': 'the Russian Federation', 'республики болгария': 'the Republic of Bulgaria',
    'болгарии': 'the Republic of Bulgaria', 'республики казахстан': 'the Republic of Kazakhstan',
    'казахстана': 'the Republic of Kazakhstan', 'республики беларусь': 'the Republic of Belarus',
    'беларуси': 'the Republic of Belarus', 'украины': 'Ukraine', 'армении': 'the Republic of Armenia',
    'узбекистана': 'the Republic of Uzbekistan', 'киргизии': 'the Kyrgyz Republic',
}


# В паспорте гражданство стоит в именительном («Российская Федерация»),
# а в договоре нужен родительный: «гражданин Российской Федерации».
CITIZENSHIP_RU = {
    'российская федерация': 'Российской Федерации', 'россия': 'Российской Федерации',
    'республика болгария': 'Республики Болгария', 'болгария': 'Республики Болгария',
    'республика казахстан': 'Республики Казахстан', 'казахстан': 'Республики Казахстан',
    'республика беларусь': 'Республики Беларусь', 'беларусь': 'Республики Беларусь',
    'украина': 'Украины', 'республика армения': 'Республики Армения', 'армения': 'Республики Армения',
    'республика узбекистан': 'Республики Узбекистан', 'узбекистан': 'Республики Узбекистан',
    'киргизская республика': 'Киргизской Республики', 'киргизия': 'Киргизской Республики',
}


def citizenship_ru(value: str) -> str:
    """Приводим к родительному. Приставку «гражданин» срезаем — в договоре
    она уже есть в шаблоне строки, иначе выйдет «гражданин гражданин ...»."""
    v = (value or '').strip()
    stripped = v
    for pref in ('гражданина', 'гражданин', 'гражданки', 'гражданка'):
        if stripped.lower().startswith(pref):
            stripped = stripped[len(pref):].strip()
            break
    return CITIZENSHIP_RU.get(stripped.lower(), stripped)


def citizenship_en(value: str) -> str:
    """Гражданство → английская колонка.

    Паспорт даёт именительный («Российская Федерация»), договоры и прошлые
    сделки — родительный («Российской Федерации»). Сначала приводим к
    родительному, потом переводим: одна таблица вместо двух.
    """
    # lstrip() режет символы, а не префикс — здесь нужен именно removeprefix
    v = (value or '').strip()
    key = v.lower().removeprefix('гражданин').removeprefix('гражданка').strip()
    genitive = CITIZENSHIP_RU.get(key, v)
    for candidate in (genitive.lower(), key, v.lower()):
        if candidate in CITIZENSHIP_EN:
            return CITIZENSHIP_EN[candidate]
    return v


def client_line(f: dict, lang: str = 'ru') -> str:
    """Преамбула клиента: ФИО, гражданство, паспорт, выдан, срок, ДР, нац. ID."""
    if lang == 'ru':
        parts = [f.get('client_name_ru') or f.get('client_name_en') or '']
        if f.get('client_name_en') and f.get('client_name_ru'):
            parts[0] += f" ({f['client_name_en']})"
        if f.get('client_citizenship'):
            parts.append(f"гражданин {citizenship_ru(f['client_citizenship'])}")
        if f.get('client_passport_no'):
            parts.append(f"паспорт № {f['client_passport_no']}")
        if f.get('client_passport_issue_date'):
            issued = f"выдан {f['client_passport_issue_date']}"
            if f.get('client_passport_issued_by'):
                issued += f" {f['client_passport_issued_by']}"
            parts.append(issued)
        if f.get('client_passport_expiry_date'):
            parts.append(f"действителен до {f['client_passport_expiry_date']}")
        if f.get('client_birth_date'):
            parts.append(f"дата рождения {f['client_birth_date']}")
        if f.get('client_national_id'):
            parts.append(f"персональный № {f['client_national_id']}")
        return ', '.join(p for p in parts if p)

    parts = [f.get('client_name_en') or f.get('client_name_ru') or '']
    if f.get('client_citizenship'):
        parts.append(f"a citizen of {citizenship_en(f['client_citizenship'])}")
    if f.get('client_passport_no'):
        parts.append(f"passport No. {f['client_passport_no']}")
    if f.get('client_passport_issue_date'):
        parts.append(f"issued on {f['client_passport_issue_date']}")
    if f.get('client_passport_expiry_date'):
        parts.append(f"valid until {f['client_passport_expiry_date']}")
    if f.get('client_birth_date'):
        parts.append(f"date of birth {f['client_birth_date']}")
    if f.get('client_national_id'):
        parts.append(f"personal No. {f['client_national_id']}")
    return ', '.join(p for p in parts if p)


def property_line(f: dict, deal_type: str) -> str:
    bits = []
    if f.get('project_name'):
        bits.append(f['project_name'])
    if f.get('unit_no'):
        bits.append(f"юнит / unit {f['unit_no']}")
    if f.get('property_address'):
        bits.append(f"адрес / address: {f['property_address']}")
    if deal_type == 'rental':
        if f.get('rent_amount'):
            bits.append(f"арендная плата / rent: {f['rent_amount']}")
        if f.get('deposit_amount'):
            bits.append(f"депозит / deposit: {f['deposit_amount']}")
    if deal_type == 'leasehold':
        if f.get('lease_term'):
            bits.append(f"срок leasehold / lease term: {f['lease_term']}")
        if f.get('lease_start_date'):
            bits.append(f"дата начала / start date: {f['lease_start_date']}")
    return ';\n'.join(bits)


def recipient_details(f: dict) -> str:
    rows = []
    if f.get('recipient_bank'):
        rows.append(f"Банк / Bank: {f['recipient_bank']}")
    if f.get('recipient_account'):
        rows.append(f"Счёт / Account No.: {f['recipient_account']}")
    if f.get('recipient_swift'):
        rows.append(f"SWIFT: {f['recipient_swift']}")
    if f.get('recipient_bik'):
        rows.append(f"БИК / BIK: {f['recipient_bik']}")
    if f.get('recipient_bank_address'):
        rows.append(f"Адрес банка / Bank address: {f['recipient_bank_address']}")
    return '\n'.join(rows)


def payment_basis(f: dict, deal_type: str) -> str:
    if deal_type == 'rental' and not f.get('invoice_no'):
        base = 'Договор аренды (Lease Agreement)'
        if f.get('contract_ref'):
            base += f" {f['contract_ref']}"
        return base + ';\nотдельный инвойс не выставлялся / no separate invoice issued'
    bits = []
    if f.get('invoice_no'):
        inv = f"Инвойс / Invoice № {f['invoice_no']}"
        if f.get('invoice_date'):
            inv += f" от / dated {f['invoice_date']}"
        bits.append(inv)
    if f.get('contract_ref'):
        bits.append(f"Договор-основание / Underlying contract: {f['contract_ref']}")
    return ';\n'.join(bits)


def building_of(unit: str) -> str:
    """Корпус из номера юнита: B2-711 → B2. У вилл и участков корпуса нет."""
    u = (unit or '').strip()
    if '-' in u:
        head = u.split('-', 1)[0].strip()
        if head and len(head) <= 4:
            return head
    return ''


def property_description(f: dict, deal_type: str) -> str:
    """Строка позиции коммерческого инвойса.

    Формат снят с реальных инвойсов MF-180-2808-1 и MF-180-3108-1:
    «Оплата по графику за апартаменты HEART BY BOTANICA (PHASE 1),
    Building B2, Unit B2-711 (leasehold).»
    """
    kind = {'freehold': 'freehold', 'leasehold': 'leasehold', 'rental': 'аренда'}.get(deal_type, '')
    bits = []
    if f.get('project_name'):
        bits.append(str(f['project_name']))
    building = building_of(f.get('unit_no', ''))
    if building:
        bits.append(f'Building {building}')
    if f.get('unit_no'):
        bits.append(f"Unit {f['unit_no']}")
    body = ', '.join(bits) if bits else 'объект по договору'
    tail = f' ({kind})' if kind else ''
    return f'Оплата по графику за апартаменты {body}{tail}.'


def payment_reference(f: dict, method: str, part: int | None = None) -> str:
    """Назначение платежа для банка клиента.

    Формулировка каноническая — переписана с реальных инвойсов MF Corp:
    «ОПЛАТА ПО ИНВОЙСУ № BBB1-2026018 ОТ 26.08.2026 ЗА АПАРТАМЕНТЫ
    UNIT B2-711, BUILDING B2, HEART BY BOTANICA (PHASE 1),
    ДЛЯ NADEZHDA BUROVA. БЕЗ НДС.»

    Капс, дата инвойса, корпус и ФИО ЛАТИНИЦЕЙ — банк сверяет платёж с
    инвойсом застройщика, поэтому отсебятина здесь дороже всего.
    """
    if method in ('sbp', 'usdt', 'cash'):
        return ('Без указания назначения платежа; подтверждение оплаты направляется менеджеру MF / '
                'No payment reference in the transfer; the payment evidence is sent to the MF manager')
    head = 'ОПЛАТА ПО ИНВОЙСУ'
    if f.get('invoice_no'):
        head += f" № {f['invoice_no']}"
    if f.get('invoice_date'):
        head += f" ОТ {f['invoice_date']}"
    obj = []
    if f.get('unit_no'):
        obj.append(f"UNIT {f['unit_no']}")
    building = building_of(f.get('unit_no', ''))
    if building:
        obj.append(f'BUILDING {building}')
    if f.get('project_name'):
        obj.append(str(f['project_name']))
    # после даты инвойса запятой нет — «… ОТ 26.08.2026 ЗА АПАРТАМЕНТЫ …»
    if obj:
        head += ' ЗА АПАРТАМЕНТЫ ' + ', '.join(obj)
    parts = [head]
    # в назначении ФИО латиницей — так его сверяет банк-получатель
    name = f.get('client_name_en') or f.get('client_name_ru') or ''
    if name:
        parts.append(f'ДЛЯ {name}')
    if part and part > 1:
        parts.append(f'ЧАСТЬ {part}')
    return (', '.join(parts) + '. БЕЗ НДС.').upper()


# ─────────────────────── заполнение приложений ───────────────────────

def _fill_client_signature(table, f: dict, money: dict) -> None:
    """Блок подписи Клиента — он повторяется в теле договора и в Приложении 1."""
    for row in table.rows:
        for cell in row.cells:
            if '[Full name / legal name]' not in cell.text:
                continue
            for p in cell.paragraphs:
                _para_replace(p, '[Full name / legal name]',
                              f.get('client_name_ru') or f.get('client_name_en') or '')
                _para_replace(p, '[ID / Registration No.]', f.get('client_passport_no') or '')
                _para_replace(p, '[Address]', money.get('client_address') or '')
                _para_replace(p, '[Email / phone]', money.get('client_contact') or '')
                _para_replace(p, '[Authorised signatory, if applicable]', '—')


def _fill_appendix1(doc, f: dict, deal_type: str, money: dict, number: str, when: datetime) -> None:
    t = doc.tables[1]
    # Формулировки дословно из живых документов: у Фролова (аренда) и Буровой
    # (лизхолд) комиссия «в курсе», у Антоненко (фрихолд) курса RUB/THB нет
    # вовсе — там «в согласованной сумме pay-in».
    if deal_type == 'freehold':
        default_fee = ('Включена в согласованную сумму pay-in; отдельно не взимается / '
                       'Included in the agreed pay-in amount; no separate charge')
    else:
        default_fee = (f"Включена в курс {money.get('rate', '')} RUB/THB, отдельно не взимается / "
                       f"Included in the rate of {money.get('rate', '')} RUB/THB, "
                       f"not charged separately")
    fee_note = money.get('fee_note') or default_fee

    _set_field(t, 'Номер и дата', f'№ {number} от {date_ru_en(when)}')
    _set_field(t, 'Клиент / Client', client_line(f, 'ru') + '\n' + client_line(f, 'en'))
    _set_field(t, 'Основание платежа', payment_basis(f, deal_type))
    _set_field(t, 'Объект / Property', property_line(f, deal_type))
    _set_field(t, 'Получатель / Recipient', f.get('recipient_name'))
    _set_field(t, 'Реквизиты получателя', recipient_details(f))
    _set_field(t, 'Комиссия Агента / Agent', fee_note)
    _set_field(t, 'Всего к оплате Клиентом', money.get('total_payin'))
    _set_field(t, 'Сумма для перечисления получателю', money.get('transfer_amount'))
    _set_field(t, 'Остаток после исполнения', DEFAULTS['remaining_balance'])
    days = money.get('execution_days') or EXECUTION_DAYS.get(deal_type, DEFAULTS['execution_days'])
    word = 'рабочий день' if str(days) == '1' else 'рабочих дня'
    _set_field(t, 'Срок исполнения',
               f'{days} {word} после выполнения п. 2.2 Договора / '
               f'{days} business day(s) after Clause 2.2 is met')
    _set_field(t, 'Контакт для подтверждения', DEFAULTS['confirmation_contact'])

    # итоговый блок: обе комиссии по умолчанию вшиты в курс
    rows = [r for r in t.rows if 'Комиссия Агента' in r.cells[0].text]
    if len(rows) > 1:
        _set_cell(rows[-1].cells[-1], DEFAULTS['fee_included'])
    _set_field(t, 'Комиссия платёжного партнёра', DEFAULTS['fee_included'])

    if deal_type == 'freehold':
        _set_field(t, 'Банковские расходы', money.get('bank_charges') or 'OUR')
        _set_field(t, 'Сумма и валюта pay-in', money.get('total_payin'))
        _set_field(t, 'Обязательство по инвойсу застройщика',
                   f"{f.get('invoice_currency') or 'THB'} {f.get('invoice_amount') or ''}".strip())
        _set_field(t, 'Источник курса и срок действия', money.get('rate_source'))
        _set_field(t, 'Подтверждённый USD-эквивалент', money.get('usd_equivalent'))
        _set_field(t, 'Статус зачёта THB-инвойса', money.get('thb_credit_status'))
        _set_field(t, 'Письменное подтверждение застройщика', money.get('developer_confirmation'))
    else:
        _set_field(t, 'Банковские расходы', DEFAULTS['bank_charges_local'])
        _set_field(t, 'Сумма, поступающая Агенту', money.get('total_payin'))
        _set_field(t, 'Сумма и валюта перевода', money.get('transfer_amount'))
        _set_field(t, 'Курс и срок его действия',
                   f"{money.get('rate', '')} RUB за 1 THB / {money.get('rate', '')} RUB per 1 THB — "
                   f"до {money.get('rate_valid_until') or f'{when:%d.%m.%Y}, 23:59 (GMT+7)'}")

    if deal_type == 'leasehold':
        _set_field(t, 'да / нет / не подтверждено', money.get('registration_needed') or 'не подтверждено / not confirmed')
        _set_field(t, 'клиент / застройщик / арендодатель',
                   money.get('registration_by') or 'застройщик / developer')
    if deal_type == 'rental':
        _set_field(t, 'Вид платежа', money.get('payment_type') or 'депозит / deposit')

    _fill_client_signature(t, f, money)


def _fill_appendix2(doc, f: dict, money: dict, number: str, when: datetime) -> None:
    """Приложение 2 — маршрут pay-in.

    Реквизиты и назначение здесь НЕ дублируются, а отсылают к коммерческому
    инвойсу: так в обоих живых договорах — у Буровой «Указывается в Invoice
    № MF-180-2808-1 / As stated in the Invoice», у Антоненко «По отдельному
    коммерческому инвойсу Агента». Дублировать опасно: реквизиты меняются,
    а подписанный договор — нет.
    """
    t = doc.tables[2]
    method = money.get('payin_method') or 'bank'
    valid = money.get('rate_valid_until') or f'{when:%d.%m.%Y}, 23:59 (GMT+7)'
    agent = AGENT['name']
    _set_field(t, 'Способ pay-in', PAYIN_METHODS.get(method, method))
    _set_field(t, 'Валюта / Currency', money.get('payin_currency') or 'RUB')
    _set_field(t, 'Получатель платежа',
               f'Указывается в актуальном Invoice № {number}: {agent} либо уполномоченный '
               f'платёжный партнёр Агента / As stated in the current Invoice: {agent} '
               f'or the Agent authorised payment partner')
    _set_field(t, 'Роль получателя',
               'Агент либо указанный им уполномоченный получатель / '
               'Agent or its authorised pay-in recipient')
    _set_field(t, 'Реквизиты / Payment details',
               money.get('payin_details')
               or f'Банковские реквизиты указываются в актуальном Invoice № {number} '
                  f'с номером, датой и сроком действия реквизитов; переносить реквизиты '
                  f'из предыдущих сделок запрещено / Bank details are stated in the current '
                  f'Invoice bearing its number, date and validity period')
    _set_field(t, 'Invoice / Instruction No.', f'№ {number} от {when:%d.%m.%Y}')
    _set_field(t, 'Срок действия реквизитов', valid)
    _set_field(t, 'Назначение платежа',
               f'Указывается в Invoice № {number} / As stated in the Invoice')
    _set_field(t, 'Подтверждение оплаты', PAYIN_EVIDENCE.get(method, ''))


# ─────────────────────────── публичное API ───────────────────────────

ANNEX_FORM_MARK = ('ФОРМА — заполняется отдельным дополнительным соглашением на каждый платёж / '
                   'TEMPLATE — completed by a separate supplementary agreement for each payment')


def build_agreement(deal_type: str, fields: dict, money: dict,
                    number: str | None = None, when: datetime | None = None,
                    blank_annexes: bool = True) -> tuple[bytes, str]:
    """Рамочный договор. → (docx-байты, номер).

    По умолчанию Приложения 1 и 2 остаются пустыми формами: договор клиент
    подписывает один раз и держит у себя неизменным, а суммы и курс каждого
    платежа живут в отдельном дополнительном соглашении. Заполненные
    приложения внутри договора сделали бы его привязанным к первому платежу.
    """
    from docx import Document  # noqa: PLC0415

    if deal_type not in TEMPLATES:
        raise ValueError(f'неизвестный тип сделки: {deal_type}')
    when = when or datetime.now()
    number = number or make_number(fields.get('client_passport_no', ''), when, 1)

    doc = Document(os.path.join(TEMPLATE_DIR, TEMPLATES[deal_type]))
    _strip_draft_mark(doc)

    for p in doc.paragraphs:
        _para_replace(p, 'г. [●] / [●], [дата / date]',
                      f'г. Пхукет, Таиланд / Phuket, Thailand, {date_ru_en(when)}')

    body = doc.tables[0]
    for row in body.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _para_replace(p, '[ФИО / полное наименование, ID/регистрационный номер, адрес]',
                              client_line(fields, 'ru'))
                _para_replace(p, '[full name / legal name, ID or registration number, address]',
                              client_line(fields, 'en'))
                _para_replace(p, '[учредительных документов / corporate documents]',
                              'учредительных документов')
                _para_replace(p, '[corporate documents]', 'the corporate documents')
                _para_replace(p, '[Full name / legal name]',
                              fields.get('client_name_ru') or fields.get('client_name_en') or '')
                _para_replace(p, '[ID / Registration No.]', fields.get('client_passport_no') or '')
                _para_replace(p, '[Address]', money.get('client_address') or '')
                _para_replace(p, '[Email / phone]', money.get('client_contact') or '')
                _para_replace(p, '[Authorised signatory, if applicable]', '—')
                _para_replace(p, '[5]', DEFAULTS['report_days'])
                _para_replace(p, '[10]', DEFAULTS['terminate_days'])

    row = _find_row(body, 'рабочих дней после выполнения пункта 2.2')
    if row is not None:
        # в rental пункт лежит в колонках 0/1, во freehold — в 1/2, поэтому все
        for cell in row.cells:
            for p in cell.paragraphs:
                _para_replace(p, '[●]', money.get('execution_days') or DEFAULTS['execution_days'])

    if blank_annexes:
        for para in doc.paragraphs:
            if para.text.strip().upper().startswith('ПРИЛОЖЕНИЕ'):
                mark = para.insert_paragraph_before(ANNEX_FORM_MARK)
                for run in mark.runs:
                    run.italic = True
        _fill_client_signature(doc.tables[1], fields, money)
    else:
        _fill_appendix1(doc, fields, deal_type, money, number, when)
        _fill_appendix2(doc, fields, money, number, when)
    _sign_agent(doc, when)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), number


def build_addendum(deal_type: str, fields: dict, money: dict, number: str,
                   parent_number: str, seq: int, when: datetime | None = None) -> bytes:
    """Доп. соглашение на очередной платёж: только Приложения 1 и 2.

    Договор не перевыпускаем — допник ссылается на него, как и требует шаблон
    инвойса («договор № … от …; платёжная инструкция № … от …»).
    """
    from docx import Document  # noqa: PLC0415

    when = when or datetime.now()
    doc = Document(os.path.join(TEMPLATE_DIR, TEMPLATES[deal_type]))
    _strip_draft_mark(doc)
    _fill_appendix1(doc, fields, deal_type, money, number, when)
    _fill_appendix2(doc, fields, money, number, when)
    _sign_agent(doc, when)

    # выкидываем тело договора — остаются только приложения
    body = doc.tables[0]._element
    body.getparent().remove(body)
    kept = 0
    for p in list(doc.paragraphs):
        if 'ПРИЛОЖЕНИЕ' in p.text.upper():
            kept += 1
        if kept == 0:
            p._element.getparent().remove(p._element)

    # insert_paragraph_before вставляет ПЕРЕД якорем, поэтому якорь один и тот же,
    # а строки идут в нужном порядке: RU → EN → место и дата
    anchor = doc.paragraphs[0]
    anchor.insert_paragraph_before(
        f'ДОПОЛНИТЕЛЬНОЕ СОГЛАШЕНИЕ № {seq} к Агентскому договору № {parent_number}')
    anchor.insert_paragraph_before(
        f'SUPPLEMENTARY AGREEMENT No. {seq} to Agency Agreement No. {parent_number}')
    anchor.insert_paragraph_before(
        f'г. Пхукет, Таиланд / Phuket, Thailand, {date_ru_en(when)}')

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_invoice(fields: dict, money: dict, number: str, parent_number: str,
                  instruction_number: str, when: datetime | None = None) -> bytes:
    """Счёт клиенту на pay-in. Всегда ссылается на рамочный договор."""
    from docx import Document  # noqa: PLC0415

    when = when or datetime.now()
    doc = Document(os.path.join(TEMPLATE_DIR, INVOICE_TEMPLATE))
    _strip_draft_mark(doc)
    method = money.get('payin_method') or 'bank'
    valid = money.get('rate_valid_until') or f'{when:%d.%m.%Y}, 23:59 (GMT+7)'
    repl = {
        '[MF-INV-●]': number,
        '[дд.мм.гггг, время]': valid,
        '[дд.мм.гггг]': f'{when:%d.%m.%Y}',
        '[ФИО / наименование компании]': fields.get('client_name_ru')
                                          or fields.get('client_name_en') or '',
        '[проект, юнит / объект]': property_line(fields, money.get('deal_type') or 'leasehold'),
        '[договор № … от …; платёжная инструкция № … от …]':
            f'Договор № {parent_number}; платёжная инструкция № {instruction_number} '
            f'от {when:%d.%m.%Y}',
        '[например: оплата по сделке / бронирование / очередной платёж]':
            money.get('payment_purpose') or 'очередной платёж по договору',
        '[включена в сумму / дополнительно: …]': 'включена в сумму',
        '[валюта]': money.get('payin_currency') or 'RUB',
        '[сумма]': str(money.get('total_payin') or ''),
        '[банковский перевод / СБП / USDT / иной согласованный способ]':
            PAYIN_METHODS.get(method, method),
        '[полное наименование получателя]': AGENT['name'],
        '[счёт, банк, БИК / СБП / сеть и адрес кошелька]': money.get('payin_details') or '',
        '[указать строго в этой формулировке]':
            money.get('payment_reference') or payment_reference(fields, method, money.get('part')),
    }
    for p in doc.paragraphs:
        for old, new in repl.items():
            _para_replace(p, old, new)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old, new in repl.items():
                        _para_replace(p, old, new)
    _sign_agent(doc, when)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()



def money_ru(amount, currency: str = 'RUB') -> str:
    """«2 800 000 руб.» — формат сумм в коммерческом инвойсе."""
    raw = str(amount or '').replace('\u00a0', ' ').replace(' ', '').replace(',', '.')
    try:
        value = float(raw)
    except ValueError:
        return str(amount or '')
    whole = f'{value:,.0f}'.replace(',', '\u00a0') if value == int(value) else \
        f'{value:,.2f}'.replace(',', '\u00a0')
    suffix = {'RUB': 'руб.', 'THB': '฿', 'USD': 'USD', 'USDT': 'USDT'}.get(currency, currency)
    return f'{whole} {suffix}'


def build_commercial_invoice(fields: dict, money: dict, number: str,
                             deal_type: str, when: datetime | None = None) -> bytes:
    """Рублёвый коммерческий инвойс клиенту — тот, что уходит в банк.

    Двуязычный Invoice Short от юриста в реальных сделках не используется:
    клиент платит по этому документу, и банк сверяет назначение платежа с ним.
    """
    from docx import Document  # noqa: PLC0415

    when = when or datetime.now()
    doc = Document(os.path.join(TEMPLATE_DIR, COMMERCIAL_INVOICE_TEMPLATE))
    method = money.get('payin_method') or 'bank'
    client_bits = [fields.get('client_name_en') or fields.get('client_name_ru') or '']
    if fields.get('client_passport_no'):
        client_bits.append(f"паспорт {fields['client_passport_no']}")
    if fields.get('client_passport_issue_date'):
        client_bits.append(f"дата выдачи {fields['client_passport_issue_date']}")
    repl = {
        '[CLIENT]': ', '.join(b for b in client_bits if b),
        '[NUMBER]': number,
        '[DATE]': f'{when:%d.%m.%Y}',
        '[DESCRIPTION]': money.get('invoice_description')
                         or property_description(fields, deal_type),
        '[AMOUNT]': money_ru(money.get('total_payin'), money.get('payin_currency') or 'RUB'),
        '[PURPOSE]': money.get('payment_reference')
                     or payment_reference(fields, method, money.get('part')),
    }
    targets = list(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                targets.extend(cell.paragraphs)
    for p in targets:
        for old, new in repl.items():
            _para_replace(p, old, new)
    _sign_agent(doc, when)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────── проверка перед выдачей ───────────────────────

PLACEHOLDER_RE = re.compile(r'\[●\]|\[[^\]\n]{2,80}?\s/\s[^\]\n]{2,80}?\]')


def check(data: bytes, allow_forms: bool = False) -> list[str]:
    """Аналог check_doc.py: не отдаём документ с незаполненными местами.

    `allow_forms` — для рамочного договора: приложения в нём намеренно пустые
    бланки, плейсхолдеры там не дефект. В допнике и инвойсе — дефект всегда.
    """
    from docx import Document  # noqa: PLC0415

    doc = Document(io.BytesIO(data))
    problems = []
    texts = [p.text for p in doc.paragraphs]
    body_only = allow_forms and len(doc.tables) > 1
    for idx, t in enumerate(doc.tables):
        if body_only and idx > 0:
            continue          # таблицы 1 и 2 — бланки приложений
        for row in t.rows:
            for cell in row.cells:
                texts.append(cell.text)
    for txt in texts:
        if txt.strip() == 'ОБЕЗЛИЧЕННЫЙ ШАБЛОН':
            problems.append('осталась пометка «ОБЕЗЛИЧЕННЫЙ ШАБЛОН»')
        for m in PLACEHOLDER_RE.findall(txt):
            problems.append(f'незаполненный плейсхолдер: {m}')
    seen, out = set(), []
    for p in problems:
        if p not in seen:
            seen.add(p)
            out.append(p)
    return out
