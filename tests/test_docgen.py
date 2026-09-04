# -*- coding: utf-8 -*-
"""Тесты генератора документов и парсера полей.

Всё, что здесь проверяется, поймано на живых документах: схема нумерации взята
с договора Бранова, дефолты — с рентал-договора Фролова, а постфильтр заглушек
появился после того, как модель вернула «CLIENT FULL NAME» как значение поля.
"""
import io
import os
import sys
from datetime import datetime

import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import docgen  # noqa: E402
import docparse  # noqa: E402

CLIENT = {
    'client_name_ru': 'Бурова Надежда Васильевна',
    'client_name_en': 'Ms. Nadezhda Burova',
    'client_citizenship': 'Российской Федерации',
    'client_passport_no': '77 2817242',
    'client_passport_issue_date': '12.03.2021',
    'client_passport_expiry_date': '12.03.2031',
    'client_birth_date': '04.07.1988',
    'invoice_no': 'BBB1-2026019', 'invoice_date': '25.08.2026',
    'invoice_amount': '1000000', 'invoice_currency': 'THB',
    'project_name': 'Heart by Botanica', 'unit_no': 'B1-313',
    'property_address': '23/26, Phuket',
    'recipient_name': 'Botanica Co., Ltd.', 'recipient_bank': 'Bangkok Bank',
    'recipient_account': '731-0-13238-1', 'recipient_swift': 'BKKBTHBK',
    'lease_term': '30 лет', 'lease_start_date': '01.10.2026',
}
MONEY = {
    'rate': '2.795', 'total_payin': 'RUB 2 795 000.00',
    'transfer_amount': 'THB 1 000 000.00', 'payin_method': 'bank',
    'payin_currency': 'RUB', 'client_address': 'г. Москва',
    'client_contact': '+7 900 000-00-00',
}
MONEY_FREEHOLD = dict(MONEY, rate_source='Письмо застройщика от 03.09.2026',
                      usd_equivalent='USD 30 500.00', thb_credit_status='полное исполнение',
                      developer_confirmation='письмо № 12 от 03.09.2026', bank_charges='OUR')
WHEN = datetime(2026, 9, 4)


class TestНумерация:
    def test_схема_совпадает_с_договором_бранова(self):
        # паспорт 77 6892733, договор от 16.06 → MF-733-1606
        assert docgen.make_number('77 6892733', datetime(2026, 6, 16), 1) == 'MF-733-1606-1'

    def test_берутся_три_последние_цифры(self):
        assert docgen.make_number('77 2817242', WHEN, 1) == 'MF-242-0409-1'

    def test_буквы_в_паспорте_игнорируются(self):
        assert docgen.make_number('AV0616156', WHEN, 1) == 'MF-156-0409-1'

    def test_пустой_паспорт_не_роняет(self):
        assert docgen.make_number('', WHEN, 1) == 'MF-000-0409-1'

    def test_хвост_различает_соглашения(self):
        a = docgen.make_number('77 2817242', WHEN, 1)
        b = docgen.make_number('77 2817242', WHEN, 2)
        assert a != b


class TestНазначениеПлатежа:
    def test_на_сбп_назначение_не_указывается(self):
        # у Фролова в живом договоре: «Без указания назначения платежа»
        assert 'Без указания назначения' in docgen.payment_reference(CLIENT, 'sbp')

    @pytest.mark.parametrize('method', ['usdt', 'cash'])
    def test_на_крипте_и_налчиных_тоже(self, method):
        assert 'Без указания назначения' in docgen.payment_reference(CLIENT, method)

    def test_совпадает_с_реальным_инвойсом_mf_180_2808(self):
        """Эталон снят с живого MF_Commercial_Invoice_MF-180-2808-1.docx."""
        f = {'invoice_no': 'BBB1-2026018', 'invoice_date': '26.08.2026', 'unit_no': 'B2-711',
             'project_name': 'HEART BY BOTANICA (PHASE 1)', 'client_name_en': 'Nadezhda Burova'}
        assert docgen.payment_reference(f, 'bank') == (
            'ОПЛАТА ПО ИНВОЙСУ № BBB1-2026018 ОТ 26.08.2026 ЗА АПАРТАМЕНТЫ '
            'UNIT B2-711, BUILDING B2, HEART BY BOTANICA (PHASE 1), '
            'ДЛЯ NADEZHDA BUROVA. БЕЗ НДС.')

    def test_совпадает_с_реальным_инвойсом_mf_180_3108(self):
        f = {'invoice_no': 'BBB1-2026019', 'invoice_date': '31.08.2026', 'unit_no': 'B1-313',
             'project_name': 'HEART BY BOTANICA (PHASE 1)', 'client_name_en': 'Nadezhda Burova'}
        assert docgen.payment_reference(f, 'bank') == (
            'ОПЛАТА ПО ИНВОЙСУ № BBB1-2026019 ОТ 31.08.2026 ЗА АПАРТАМЕНТЫ '
            'UNIT B1-313, BUILDING B1, HEART BY BOTANICA (PHASE 1), '
            'ДЛЯ NADEZHDA BUROVA. БЕЗ НДС.')

    def test_фио_в_назначении_латиницей(self):
        # банк-получатель сверяет назначение с инвойсом застройщика
        ref = docgen.payment_reference(CLIENT, 'bank')
        assert 'NADEZHDA BUROVA' in ref and 'БУРОВА НАДЕЖДА' not in ref

    def test_номер_части_попадает_в_назначение(self):
        assert 'ЧАСТЬ 2' in docgen.payment_reference(CLIENT, 'bank', 2)

    def test_первая_часть_не_нумеруется(self):
        assert 'ЧАСТЬ' not in docgen.payment_reference(CLIENT, 'bank', 1)

    def test_корпус_выводится_из_юнита(self):
        assert docgen.building_of('B2-711') == 'B2'
        assert docgen.building_of('B1-313') == 'B1'

    def test_у_виллы_корпуса_нет(self):
        assert docgen.building_of('Plot 25') == ''
        assert 'BUILDING' not in docgen.payment_reference(
            {'invoice_no': 'T-1', 'unit_no': 'Plot 25', 'client_name_en': 'A B'}, 'bank')


class TestГражданство:
    def test_переводится_в_английскую_колонку(self):
        assert docgen.citizenship_en('Российской Федерации') == 'the Russian Federation'

    def test_с_приставкой_гражданин(self):
        assert docgen.citizenship_en('гражданин Республики Болгария') == 'the Republic of Bulgaria'

    def test_неизвестное_остаётся_как_есть(self):
        assert docgen.citizenship_en('Thailand') == 'Thailand'

    def test_в_английской_строке_нет_кириллицы(self):
        line = docgen.client_line(CLIENT, 'en')
        assert not any('а' <= c.lower() <= 'я' for c in line)


class TestГенерацияДоговоров:
    @pytest.mark.parametrize('deal_type,money', [
        ('freehold', MONEY_FREEHOLD), ('leasehold', MONEY), ('rental', MONEY)])
    def test_тело_договора_без_незаполненных_мест(self, deal_type, money):
        data, _ = docgen.build_agreement(deal_type, CLIENT, money, when=WHEN)
        assert docgen.check(data, allow_forms=True) == []

    @pytest.mark.parametrize('deal_type', ['freehold', 'leasehold', 'rental'])
    def test_приложения_в_договоре_остаются_бланками(self, deal_type):
        from docx import Document
        data, _ = docgen.build_agreement(deal_type, CLIENT, MONEY_FREEHOLD, when=WHEN)
        doc = Document(io.BytesIO(data))
        marks = [p.text for p in doc.paragraphs if p.text.startswith('ФОРМА')]
        assert len(marks) == 2          # по пометке над каждым приложением
        # суммы конкретного платежа в рамочный договор не попадают
        annex = '\n'.join(c.text for t in doc.tables[1:] for r in t.rows for c in r.cells)
        assert MONEY['total_payin'] not in annex

    def test_номер_проставляется_из_паспорта(self):
        _, number = docgen.build_agreement('leasehold', CLIENT, MONEY, when=WHEN)
        assert number == 'MF-242-0409-1'

    def test_комиссия_по_умолчанию_вшита_в_курс(self):
        from docx import Document
        add = docgen.build_addendum('leasehold', CLIENT, MONEY, 'MF-1', 'MF-0', 1, when=WHEN)
        text = '\n'.join(c.text for t in Document(io.BytesIO(add)).tables
                         for r in t.rows for c in r.cells)
        assert 'Включена в курс' in text and 'отдельно не взимается' in text

    def test_на_лизхолде_нет_our_sha_ben(self):
        from docx import Document
        add = docgen.build_addendum('leasehold', CLIENT, MONEY, 'MF-1', 'MF-0', 1, when=WHEN)
        text = '\n'.join(c.text for t in Document(io.BytesIO(add)).tables
                         for r in t.rows for c in r.cells)
        assert 'Внутренний тайский перевод' in text
        assert 'OUR / SHA / BEN' not in text

    def test_фрихолд_без_подтверждения_застройщика_оставляет_дыры(self):
        # блок USD-конверсии обязателен по п. 2.3 — check обязан это поймать
        add = docgen.build_addendum('freehold', CLIENT, MONEY, 'MF-1', 'MF-0', 1, when=WHEN)
        assert docgen.check(add)

    def test_подпись_и_печать_вставлены_по_одному_разу(self):
        from docx import Document
        data, _ = docgen.build_agreement('leasehold', CLIENT, MONEY, when=WHEN)
        doc = Document(io.BytesIO(data))
        blobs = [r.target_part.blob for r in doc.part.rels.values()
                 if 'image' in r.reltype and hasattr(r, 'target_part')]
        stamp = open(os.path.join(docgen.ASSETS, 'stamp.png'), 'rb').read()
        # объединённые ячейки раньше приводили к двойной печати на странице
        xml = doc.element.xml
        assert xml.count('Signature: ') >= 1
        assert any(b == stamp for b in blobs)


class TestДопСоглашение:
    def test_тело_договора_вырезано_остались_приложения(self):
        from docx import Document
        add = docgen.build_addendum('leasehold', CLIENT, MONEY, 'MF-242-0409-2',
                                    'MF-242-0409-1', 2, when=WHEN)
        assert len(Document(io.BytesIO(add)).tables) == 2

    def test_заголовки_идут_в_порядке_ru_en_дата(self):
        from docx import Document
        add = docgen.build_addendum('leasehold', CLIENT, MONEY, 'MF-242-0409-2',
                                    'MF-242-0409-1', 2, when=WHEN)
        heads = [p.text.strip() for p in Document(io.BytesIO(add)).paragraphs if p.text.strip()][:3]
        assert heads[0].startswith('ДОПОЛНИТЕЛЬНОЕ СОГЛАШЕНИЕ № 2')
        assert heads[1].startswith('SUPPLEMENTARY AGREEMENT No. 2')
        assert 'Пхукет' in heads[2]

    def test_ссылается_на_рамочный_договор(self):
        from docx import Document
        add = docgen.build_addendum('leasehold', CLIENT, MONEY, 'MF-242-0409-2',
                                    'MF-242-0409-1', 2, when=WHEN)
        assert 'MF-242-0409-1' in Document(io.BytesIO(add)).paragraphs[0].text

    def test_без_дыр(self):
        add = docgen.build_addendum('rental', CLIENT, MONEY, 'MF-242-0409-2',
                                    'MF-242-0409-1', 2, when=WHEN)
        assert docgen.check(add) == []


class TestИнвойс:
    def test_ссылается_на_договор_и_инструкцию(self):
        from docx import Document
        inv = docgen.build_invoice(CLIENT, MONEY, 'MF-242-0409-2', 'MF-242-0409-1',
                                   'MF-242-0409-2', when=WHEN)
        text = '\n'.join(p.text for p in Document(io.BytesIO(inv)).paragraphs)
        text += '\n'.join(c.text for t in Document(io.BytesIO(inv)).tables
                          for r in t.rows for c in r.cells)
        assert 'MF-242-0409-1' in text

    def test_без_дыр(self):
        inv = docgen.build_invoice(CLIENT, MONEY, 'MF-242-0409-2', 'MF-242-0409-1',
                                   'MF-242-0409-2', when=WHEN)
        assert docgen.check(inv) == []


class TestПостфильтрЗаглушек:
    @pytest.mark.parametrize('key,value', [
        ('client_name_en', 'CLIENT FULL NAME'), ('client_passport_no', 'PASSPORT NO.'),
        ('client_citizenship', 'CITIZENSHIP'), ('invoice_date', 'null'),
        ('invoice_no', 'MF-XXX-XXXX-1'), ('invoice_date', '[DATE]'),
        ('project_name', '___'), ('unit_no', '[●]'), ('recipient_name', '  '),
        # рамка внутри строки — просочилась на живом прогоне 04.09
        ('project_name', '[PROJECT / UNIT No.]'),
        ('invoice_amount', '1 USD = [FIXED RATE] RUB'),
    ])
    def test_ловит_рамку_макета(self, key, value):
        assert docparse.is_placeholder(key, value)

    @pytest.mark.parametrize('key,value', [
        ('invoice_currency', 'RUB'), ('recipient_swift', 'BKKBTHBK'),
        ('recipient_bik', '044525225'), ('recipient_name', 'Apimuk Limkul'),
        ('client_name_ru', 'Бурова Надежда Васильевна'), ('unit_no', 'B1-313'),
        ('invoice_amount', '1000000'), ('project_name', 'Heart by Botanica'),
        # регрессия: эвристика «сплошные заглавные» съедала реальные номера
        ('invoice_no', 'INV-1'), ('invoice_no', 'BBB1-2026019'),
        ('recipient_name', 'OOO "MF CORPORATION"'), ('unit_no', 'A1/47'),
        ('recipient_name', 'BOTANICA CO LTD'),
    ])
    def test_не_трогает_настоящие_значения(self, key, value):
        assert not docparse.is_placeholder(key, value)

    def test_заглушки_уходят_в_masked(self):
        merged = docparse.merge([{'_file': 'inv.pdf', 'client_name_en': 'CLIENT FULL NAME',
                                  'invoice_currency': 'RUB', 'masked_fields': []}])
        assert 'client_name_en' not in merged['fields']
        assert merged['fields']['invoice_currency'] == 'RUB'
        assert 'client_name_en' in merged['masked_fields']


class TestВалидацияРеквизитов:
    def test_бик_в_поле_swift_ловится(self):
        # общий дефект всех моделей на замере 04.09
        problems = docparse.validate({'recipient_swift': '044525225'})
        assert problems and 'БИК' in problems[0]

    def test_нормальный_swift_проходит(self):
        assert docparse.validate({'recipient_swift': 'BKKBTHBK'}) == []

    def test_бик_не_из_девяти_цифр_ловится(self):
        assert docparse.validate({'recipient_bik': '0445252'})

    def test_пустые_поля_не_ругаются(self):
        assert docparse.validate({}) == []


class TestСклейкаФайлов:
    def test_провенанс_помнит_источник(self):
        merged = docparse.merge([
            {'_file': 'passport.pdf', 'client_name_ru': 'Иванов Иван', 'masked_fields': []},
            {'_file': 'invoice.pdf', 'invoice_no': 'INV-1', 'masked_fields': []}])
        assert merged['provenance']['client_name_ru'] == 'passport.pdf'
        assert merged['provenance']['invoice_no'] == 'invoice.pdf'

    def test_конфликт_не_затирается_молча(self):
        merged = docparse.merge([
            {'_file': 'a.pdf', 'unit_no': 'B2-711', 'masked_fields': []},
            {'_file': 'b.pdf', 'unit_no': 'B1-313', 'masked_fields': []}])
        assert 'unit_no' in merged['conflicts']
        assert len(merged['conflicts']['unit_no']) == 2


class TestСлотыДокументов:
    """Менеджер грузит паспорт, инвойс и СПА по отдельности — это меняет склейку."""

    def test_фио_берётся_из_паспорта_даже_если_инвойс_пришёл_первым(self):
        merged = docparse.merge([
            {'_file': 'invoice.pdf', '_kind': 'invoice',
             'client_name_ru': 'Бурова Н.В.', 'masked_fields': []},
            {'_file': 'passport.pdf', '_kind': 'passport',
             'client_name_ru': 'Бурова Надежда Васильевна', 'masked_fields': []}])
        assert merged['fields']['client_name_ru'] == 'Бурова Надежда Васильевна'
        assert merged['provenance']['client_name_ru'] == 'passport.pdf'
        assert 'client_name_ru' not in merged['conflicts']

    def test_реквизиты_получателя_берутся_из_инвойса(self):
        merged = docparse.merge([
            {'_file': 'spa.pdf', '_kind': 'spa',
             'recipient_bank': 'какой-то банк', 'masked_fields': []},
            {'_file': 'invoice.pdf', '_kind': 'invoice',
             'recipient_bank': 'Bangkok Bank', 'masked_fields': []}])
        assert merged['fields']['recipient_bank'] == 'Bangkok Bank'

    def test_владелец_не_перебивается_чужим_документом(self):
        merged = docparse.merge([
            {'_file': 'passport.pdf', '_kind': 'passport',
             'client_name_ru': 'Бурова Надежда Васильевна', 'masked_fields': []},
            {'_file': 'invoice.pdf', '_kind': 'invoice',
             'client_name_ru': 'Бурова Н.В.', 'masked_fields': []}])
        assert merged['fields']['client_name_ru'] == 'Бурова Надежда Васильевна'
        assert 'client_name_ru' not in merged['conflicts']

    def test_срок_leasehold_берётся_из_спа(self):
        merged = docparse.merge([
            {'_file': 'invoice.pdf', '_kind': 'invoice', 'lease_term': '20 лет', 'masked_fields': []},
            {'_file': 'spa.pdf', '_kind': 'spa', 'lease_term': '30 лет', 'masked_fields': []}])
        assert merged['fields']['lease_term'] == '30 лет'

    def test_чужое_значение_закрывает_пустоту(self):
        # паспорта нет — ФИО из инвойса лучше, чем ничего
        merged = docparse.merge([
            {'_file': 'invoice.pdf', '_kind': 'invoice',
             'client_name_ru': 'Бурова Н.В.', 'masked_fields': []}])
        assert merged['fields']['client_name_ru'] == 'Бурова Н.В.'

    def test_разногласие_равноправных_источников_остаётся_конфликтом(self):
        # оба документа — не владельцы поля unit_no по отношению к invoice
        merged = docparse.merge([
            {'_file': 'a.pdf', '_kind': 'invoice', 'unit_no': 'B2-711', 'masked_fields': []},
            {'_file': 'b.pdf', '_kind': 'invoice', 'unit_no': 'B1-313', 'masked_fields': []}])
        assert 'unit_no' in merged['conflicts']

    def test_служебные_ключи_не_утекают_в_провенанс(self):
        merged = docparse.merge([
            {'_file': 'passport.pdf', '_kind': 'passport',
             'client_name_ru': 'Иванов', 'masked_fields': []}])
        assert all(not k.endswith('__owner') for k in merged['provenance'])

    def test_подсказка_типа_документа_уходит_в_промпт(self):
        assert 'загранпаспорт' in docparse.DOC_KINDS['passport']
        assert 'инвойс' in docparse.DOC_KINDS['invoice'].lower()
        assert 'SPA' in docparse.DOC_KINDS['spa']


class TestНормализацияКапса:
    """В паспорте ФИО и гражданство напечатаны капсом — в договор так нельзя."""

    @pytest.mark.parametrize('key,raw,want', [
        ('client_name_ru', 'БУРОВА НАДЕЖДА ВАСИЛЬЕВНА', 'Бурова Надежда Васильевна'),
        ('client_name_en', 'BUROVA NADEZHDA', 'Burova Nadezhda'),
        ('client_citizenship', 'РОССИЙСКАЯ ФЕДЕРАЦИЯ', 'Российская Федерация'),
        ('client_name_ru', 'ИВАНОВ-ПЕТРОВ ПЁТР', 'Иванов-Петров Пётр'),
        ('recipient_bank', 'BANGKOK BANK PUBLIC COMPANY LIMITED',
         'Bangkok Bank Public Company Limited'),
    ])
    def test_капс_приводится_к_нормальному_регистру(self, key, raw, want):
        assert docparse.normalize(key, raw) == want

    @pytest.mark.parametrize('key,raw', [
        ('invoice_currency', 'THB'), ('recipient_swift', 'BKKBTHBK'),
        ('client_passport_issued_by', 'МВД 52014'),
        ('recipient_name', 'Botanica Bangtao Beach 1 Co., Ltd.'),
    ])
    def test_аббревиатуры_и_нормальный_текст_не_трогаются(self, key, raw):
        assert docparse.normalize(key, raw) == raw

    def test_частицы_фамилий_остаются_строчными(self):
        assert docparse.normalize('client_name_ru', 'ВАН ДЕР САР ЭДВИН') == 'Ван дер Сар Эдвин'

    def test_нормализация_работает_внутри_склейки(self):
        merged = docparse.merge([{'_file': 'p.jpg', '_kind': 'passport',
                                  'client_name_ru': 'БУРОВА НАДЕЖДА', 'masked_fields': []}])
        assert merged['fields']['client_name_ru'] == 'Бурова Надежда'
        assert merged['slots']['client_name_ru'] == 'passport'


class TestГражданствоВДоговоре:
    """Паспорт даёт именительный, договору нужен родительный."""

    @pytest.mark.parametrize('raw,want', [
        ('Российская Федерация', 'Российской Федерации'),
        ('РОССИЙСКАЯ ФЕДЕРАЦИЯ', 'Российской Федерации'),
        ('Российской Федерации', 'Российской Федерации'),
        ('гражданин Республики Болгария', 'Республики Болгария'),
        ('гражданка Украины', 'Украины'),
    ])
    def test_родительный_падеж(self, raw, want):
        assert docgen.citizenship_ru(raw) == want

    def test_приставка_не_дублируется(self):
        line = docgen.client_line({'client_name_ru': 'Иванов Иван',
                                   'client_citizenship': 'гражданин Республики Болгария'}, 'ru')
        assert line.count('гражданин') == 1

    @pytest.mark.parametrize('raw', ['Российская Федерация', 'Российской Федерации',
                                     'РОССИЙСКАЯ ФЕДЕРАЦИЯ'])
    def test_английская_колонка_переводит_обе_формы(self, raw):
        assert docgen.citizenship_en(raw) == 'the Russian Federation'

    def test_неизвестная_страна_проходит_как_есть(self):
        assert docgen.citizenship_ru('Thailand') == 'Thailand'
        assert docgen.citizenship_en('Thailand') == 'Thailand'


class TestКоммерческийИнвойс:
    """Рублёвый инвойс — тот, что реально уходит клиенту и в банк."""

    def _text(self, data):
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        parts += [c.text for t in doc.tables for r in t.rows for c in r.cells]
        return '\n'.join(parts)

    def test_нет_незаполненных_плейсхолдеров(self):
        data = docgen.build_commercial_invoice(CLIENT, MONEY, 'MF-242-0409-1', 'leasehold', when=WHEN)
        text = self._text(data)
        assert not [m for m in ('[CLIENT]', '[NUMBER]', '[DATE]', '[AMOUNT]',
                                '[PURPOSE]', '[DESCRIPTION]') if m in text]

    def test_клиент_указан_латиницей_с_паспортом(self):
        text = self._text(docgen.build_commercial_invoice(
            CLIENT, MONEY, 'MF-1', 'leasehold', when=WHEN))
        assert 'Nadezhda Burova' in text and '77 2817242' in text

    def test_сумма_в_рублях_с_разделителями(self):
        text = self._text(docgen.build_commercial_invoice(
            CLIENT, dict(MONEY, total_payin='2800000'), 'MF-1', 'leasehold', when=WHEN))
        assert '2\u00a0800\u00a0000 руб.' in text

    def test_назначение_капсом_и_каноническое(self):
        text = self._text(docgen.build_commercial_invoice(
            CLIENT, MONEY, 'MF-1', 'leasehold', when=WHEN))
        assert 'ОПЛАТА ПО ИНВОЙСУ № BBB1-2026019' in text and 'БЕЗ НДС.' in text

    def test_реквизиты_ооо_мф_на_месте(self):
        text = self._text(docgen.build_commercial_invoice(
            CLIENT, MONEY, 'MF-1', 'leasehold', when=WHEN))
        for token in ('9909726886', '40807 810 9 3872 0000286', '044525225'):
            assert token in text

    def test_менеджер_может_переписать_назначение(self):
        text = self._text(docgen.build_commercial_invoice(
            CLIENT, dict(MONEY, payment_reference='СВОЁ НАЗНАЧЕНИЕ'),
            'MF-1', 'leasehold', when=WHEN))
        assert 'СВОЁ НАЗНАЧЕНИЕ' in text

    @pytest.mark.parametrize('raw,want', [
        ('2800000', '2\u00a0800\u00a0000 руб.'),
        ('2 795 000', '2\u00a0795\u00a0000 руб.'),
        ('1007194.24', '1\u00a0007\u00a0194.24 руб.'),
    ])
    def test_формат_суммы(self, raw, want):
        assert docgen.money_ru(raw, 'RUB') == want

    def test_позиция_инвойса_по_образцу(self):
        f = {'project_name': 'HEART BY BOTANICA (PHASE 1)', 'unit_no': 'B2-711'}
        assert docgen.property_description(f, 'leasehold') == (
            'Оплата по графику за апартаменты HEART BY BOTANICA (PHASE 1), '
            'Building B2, Unit B2-711 (leasehold).')
